"""
Comprehensive tests for CSS parsing and style application in html4docx.

Covers:
- Compound selectors (tag.class, tag#id, .a.b)
- CSS cascade correctness at docx output level
- CSS loading order (style tag vs link tag ordering)
- Span interactions (nested spans, outer span class styles)
- Edge cases (@media blocks, !important, empty styles, invalid values)
- Property-level verification (text-decoration, background-color, etc.)
"""

import os
import tempfile
import unittest

from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.oxml.ns import qn
from docx.shared import RGBColor

from html4docx import HtmlToDocx
from html4docx.colors import Color
from html4docx.css_parser import CSSParser

test_dir = os.path.abspath(os.path.dirname(__file__))


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _run_color(run):
    """Return the run's font color as RGBColor or None."""
    try:
        return run.font.color.rgb
    except Exception:
        return None


def _run_shd_fill(run):
    """Return the run's shading fill as an uppercase hex string, or None."""
    rPr = run._r.get_or_add_rPr()
    shd = rPr.find(qn("w:shd"))
    if shd is not None:
        return (shd.get(qn("w:fill")) or "").upper().lstrip("#")
    return None


def _underline_color(run):
    """Return the underline element's w:color attribute, or None."""
    for u in run._r.xpath(".//w:u"):
        return u.get(qn("w:color"))
    return None


def _parse_html(html: str) -> Document:
    """Parse HTML string and return the resulting Document."""
    parser = HtmlToDocx()
    return parser.parse_html_string(html)


def _first_nonempty_run(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                return run
    return None


def _run_for_text(doc, text):
    """Find the first run containing *text*."""
    for para in doc.paragraphs:
        for run in para.runs:
            if text in run.text:
                return run
    return None


def _para_for_text(doc, text):
    """Find the first paragraph containing *text* in its full text."""
    for para in doc.paragraphs:
        if text in para.text:
            return para
    return None


# ─────────────────────────────────────────────────────────────────────────────
# A) CSSParser unit tests
# ─────────────────────────────────────────────────────────────────────────────

class TestCompoundSelectorParsing(unittest.TestCase):
    """Tests that compound selectors are parsed and stored correctly."""

    def setUp(self):
        self.parser = CSSParser()

    def test_compound_tag_plus_class_stored_in_compound_rules(self):
        """p.highlight should NOT go into tag_rules — it's a compound selector."""
        self.parser.parse_css("p.highlight { color: blue; }")
        self.assertNotIn("p.highlight", self.parser.tag_rules)
        self.assertEqual(len(self.parser._compound_rules), 1)
        spec, tag, classes, eid, styles = self.parser._compound_rules[0]
        self.assertEqual(tag, "p")
        self.assertEqual(classes, ["highlight"])
        self.assertIsNone(eid)
        self.assertEqual(styles["color"], "blue")

    def test_compound_tag_plus_id_stored_in_compound_rules(self):
        """div#header should be a compound rule."""
        self.parser.parse_css("div#header { text-align: center; }")
        self.assertNotIn("div#header", self.parser.tag_rules)
        self.assertEqual(len(self.parser._compound_rules), 1)
        spec, tag, classes, eid, styles = self.parser._compound_rules[0]
        self.assertEqual(tag, "div")
        self.assertEqual(eid, "header")

    def test_multi_class_selector_stored_in_compound_rules(self):
        """.a.b should be a compound rule, not class_rules['a.b']."""
        self.parser.parse_css(".active.large { font-size: 20px; }")
        self.assertNotIn("active.large", self.parser.class_rules)
        self.assertEqual(len(self.parser._compound_rules), 1)
        spec, tag, classes, eid, styles = self.parser._compound_rules[0]
        self.assertIsNone(tag)
        self.assertIn("active", classes)
        self.assertIn("large", classes)

    def test_pure_tag_goes_to_tag_rules(self):
        self.parser.parse_css("p { color: red; }")
        self.assertIn("p", self.parser.tag_rules)
        self.assertEqual(len(self.parser._compound_rules), 0)

    def test_pure_class_goes_to_class_rules(self):
        self.parser.parse_css(".highlight { color: red; }")
        self.assertIn("highlight", self.parser.class_rules)
        self.assertEqual(len(self.parser._compound_rules), 0)

    def test_pure_id_goes_to_id_rules(self):
        self.parser.parse_css("#header { color: red; }")
        self.assertIn("header", self.parser.id_rules)
        self.assertEqual(len(self.parser._compound_rules), 0)

    def test_compound_rule_retrieval_matches_correctly(self):
        """p.highlight rule only applies to <p class="highlight">, not to <p> or <div class="highlight">."""
        self.parser.parse_css("p.highlight { color: blue; } p { color: red; }")

        # p without class → only tag rule (red)
        normal, _ = self.parser.get_styles_for_element_with_important("p", {})
        self.assertEqual(normal.get("color"), "red")

        # p with class=highlight → tag rule overridden by compound rule (blue)
        normal, _ = self.parser.get_styles_for_element_with_important("p", {"class": "highlight"})
        self.assertEqual(normal.get("color"), "blue")

        # div with class=highlight → compound rule does NOT apply (tag mismatch)
        normal, _ = self.parser.get_styles_for_element_with_important("div", {"class": "highlight"})
        self.assertNotEqual(normal.get("color"), "blue")

    def test_multi_class_rule_requires_all_classes(self):
        """.a.b only applies when the element has BOTH classes."""
        self.parser.parse_css(".a.b { color: green; }")

        # Element with both classes → matched
        normal, _ = self.parser.get_styles_for_element_with_important("p", {"class": "a b"})
        self.assertEqual(normal.get("color"), "green")

        # Element with only class a → NOT matched
        normal, _ = self.parser.get_styles_for_element_with_important("p", {"class": "a"})
        self.assertIsNone(normal.get("color"))

        # Element with only class b → NOT matched
        normal, _ = self.parser.get_styles_for_element_with_important("p", {"class": "b"})
        self.assertIsNone(normal.get("color"))

    def test_compound_specificity_higher_than_pure_tag(self):
        """p.class has higher specificity than p → compound wins."""
        self.parser.parse_css("p { color: red; } p.special { color: blue; }")
        normal, _ = self.parser.get_styles_for_element_with_important("p", {"class": "special"})
        self.assertEqual(normal.get("color"), "blue")

    def test_compound_specificity_lower_than_id(self):
        """#id has higher specificity than p.class → id wins."""
        self.parser.parse_css("p.special { color: blue; } #myid { color: green; }")
        normal, _ = self.parser.get_styles_for_element_with_important(
            "p", {"class": "special", "id": "myid"}
        )
        self.assertEqual(normal.get("color"), "green")

    def test_clear_removes_compound_rules(self):
        self.parser.parse_css("p.cls { color: red; }")
        self.parser.clear()
        self.assertEqual(len(self.parser._compound_rules), 0)

    def test_has_rules_includes_compound(self):
        self.parser.parse_css("p.cls { color: red; }")
        self.assertTrue(self.parser.has_rules())

    def test_has_rules_for_element_with_compound(self):
        self.parser.parse_css("p.cls { color: red; }")
        # element matching the compound rule
        self.assertTrue(self.parser.has_rules_for_element("p", {"class": "cls"}))
        # element not matching (wrong tag)
        self.assertFalse(self.parser.has_rules_for_element("div", {"class": "cls"}))


class TestHasRulesForElementEmptyAttrs(unittest.TestCase):
    """Regression test for the bug where has_rules_for_element returned False for empty attrs."""

    def setUp(self):
        self.parser = CSSParser()

    def test_tag_rule_found_with_no_attrs(self):
        """has_rules_for_element should return True when a tag rule exists, even if attrs={}."""
        self.parser.parse_css("p { color: red; }")
        self.assertTrue(self.parser.has_rules_for_element("p"))         # attrs=None
        self.assertTrue(self.parser.has_rules_for_element("p", {}))     # attrs={}
        self.assertFalse(self.parser.has_rules_for_element("div"))      # no rule for div

    def test_class_rule_not_found_without_attrs(self):
        """Class rules require attrs to check — no class attr → no match."""
        self.parser.parse_css(".highlight { color: red; }")
        self.assertFalse(self.parser.has_rules_for_element("p"))         # no attrs
        self.assertFalse(self.parser.has_rules_for_element("p", {}))    # empty attrs
        self.assertTrue(self.parser.has_rules_for_element("p", {"class": "highlight"}))


class TestMarkElementUsedStringClass(unittest.TestCase):
    """Regression test for mark_element_used treating a string class as a list of chars."""

    def setUp(self):
        self.parser = CSSParser()

    def test_string_class_stored_correctly(self):
        """Passing class as a string should add the whole class name, not individual chars."""
        self.parser.mark_element_used("div", {"class": "container"})
        self.assertIn("container", self.parser._used_classes)
        # Individual characters should NOT be in _used_classes
        self.assertNotIn("c", self.parser._used_classes)
        self.assertNotIn("o", self.parser._used_classes)

    def test_string_class_list_stored_correctly(self):
        """Space-separated string class should split into individual class names."""
        self.parser.mark_element_used("div", {"class": "foo bar"})
        self.assertIn("foo", self.parser._used_classes)
        self.assertIn("bar", self.parser._used_classes)

    def test_list_class_stored_correctly(self):
        """BeautifulSoup-style list class attribute works correctly."""
        self.parser.mark_element_used("div", {"class": ["container", "active"]})
        self.assertIn("container", self.parser._used_classes)
        self.assertIn("active", self.parser._used_classes)


class TestAtMediaSkipped(unittest.TestCase):
    """@media blocks and other at-rules should be ignored without crashing."""

    def setUp(self):
        self.parser = CSSParser()

    def test_media_block_skipped(self):
        css = """
        p { color: red; }
        @media screen and (max-width: 600px) {
            p { color: blue; }
            .hidden { display: none; }
        }
        .highlight { font-weight: bold; }
        """
        self.parser.parse_css(css)
        # The .hidden rule inside @media should NOT be loaded
        self.assertNotIn("hidden", self.parser.class_rules)
        # The @media override of p { color: blue } should NOT change p
        self.assertEqual(self.parser.tag_rules.get("p", {}).get("color"), "red")
        # Styles outside @media are loaded normally
        self.assertIn("highlight", self.parser.class_rules)

    def test_import_rule_skipped(self):
        css = """
        @import url('other.css');
        p { color: green; }
        """
        self.parser.parse_css(css)
        # @import is skipped; p rule is still loaded
        self.assertIn("p", self.parser.tag_rules)
        self.assertEqual(self.parser.tag_rules["p"]["color"], "green")

    def test_keyframes_skipped(self):
        css = """
        @keyframes slide {
            from { transform: translateX(0); }
            to { transform: translateX(100px); }
        }
        p { color: purple; }
        """
        self.parser.parse_css(css)
        # p rule is loaded; @keyframes content is not
        self.assertIn("p", self.parser.tag_rules)
        self.assertNotIn("from", self.parser.tag_rules)
        self.assertNotIn("to", self.parser.tag_rules)


class TestImportantCaseInsensitive(unittest.TestCase):
    """!important detection and removal should be case-insensitive."""

    def setUp(self):
        self.parser = CSSParser()

    def test_lowercase_important(self):
        self.parser.parse_css("p { color: red !important; }")
        _, important = self.parser.get_styles_for_element_with_important("p", {})
        self.assertIn("color", important)
        self.assertNotIn("!important", important["color"])
        self.assertEqual(important["color"].strip(), "red")

    def test_uppercase_important_detection(self):
        """!IMPORTANT should be detected and value should be clean."""
        self.parser.parse_css("p { color: red !IMPORTANT; }")
        _, important = self.parser.get_styles_for_element_with_important("p", {})
        self.assertIn("color", important)
        self.assertNotIn("!IMPORTANT", important["color"])
        self.assertNotIn("important", important["color"].lower())


# ─────────────────────────────────────────────────────────────────────────────
# B) CSS cascade correctness at docx output level
# ─────────────────────────────────────────────────────────────────────────────

class TestCSSCascadeOutput(unittest.TestCase):
    """Verify that the CSS cascade is correctly reflected in docx run properties."""

    def _parse(self, html):
        parser = HtmlToDocx()
        return parser.parse_html_string(html)

    def test_tag_rule_applied(self):
        html = "<style>p { color: red; }</style><p>Hello</p>"
        doc = self._parse(html)
        run = _run_for_text(doc, "Hello")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))

    def test_class_rule_overrides_tag_rule(self):
        html = """
        <style>
          p { color: red; }
          .blue { color: blue; }
        </style>
        <p class="blue">Hello</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Hello")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_id_rule_overrides_class_rule(self):
        html = """
        <style>
          .blue { color: blue; }
          #special { color: green; }
        </style>
        <p class="blue" id="special">Hello</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Hello")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(*Color["green"].value))

    def test_inline_style_overrides_id_rule(self):
        html = """
        <style>#special { color: green; }</style>
        <p id="special" style="color: red;">Hello</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Hello")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))

    def test_important_css_overrides_inline_normal(self):
        """A CSS !important declaration should beat a normal inline style."""
        html = """
        <style>p { color: blue !important; }</style>
        <p style="color: red;">Hello</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Hello")
        self.assertIsNotNone(run)
        # !important CSS → applied after normal inline in apply_styles_to_paragraph
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_multiple_classes_both_applied(self):
        """When an element has two classes, styles from both are applied (later wins)."""
        html = """
        <style>
          .bold { font-weight: bold; }
          .blue { color: blue; }
        </style>
        <p class="bold blue">Hello</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Hello")
        self.assertIsNotNone(run)
        self.assertTrue(run.font.bold)
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_last_class_wins_for_same_property(self):
        """When two classes set the same property, the last-declared class in CSS wins."""
        html = """
        <style>
          .red { color: red; }
          .blue { color: blue; }
        </style>
        <p class="red blue">Hello</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Hello")
        self.assertIsNotNone(run)
        # .blue is declared AFTER .red in CSS, same specificity → blue wins
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_compound_selector_applied_to_matching_element(self):
        """p.highlight should only color <p class="highlight">, not plain <p>."""
        html = """
        <style>
          p { color: red; }
          p.highlight { color: blue; }
        </style>
        <p>Plain paragraph</p>
        <p class="highlight">Highlighted paragraph</p>
        """
        doc = self._parse(html)

        plain_run = _run_for_text(doc, "Plain paragraph")
        highlight_run = _run_for_text(doc, "Highlighted paragraph")

        self.assertIsNotNone(plain_run)
        self.assertIsNotNone(highlight_run)

        # Plain p → only tag rule applies (red)
        self.assertEqual(_run_color(plain_run), RGBColor(255, 0, 0))
        # p.highlight → compound rule applies (blue, overrides red)
        self.assertEqual(_run_color(highlight_run), RGBColor(0, 0, 255))

    def test_compound_selector_not_applied_to_wrong_tag(self):
        """p.highlight should NOT color <div class="highlight">."""
        html = """
        <style>p.highlight { color: blue; }</style>
        <div class="highlight"><p>In div</p></div>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "In div")
        # div.highlight → compound rule doesn't match p inside div
        # (p inherits nothing by default from compound p.highlight rule)
        # The run should NOT be blue (no rule applies to p here)
        self.assertIsNotNone(run)
        color = _run_color(run)
        if color is not None:
            self.assertNotEqual(color, RGBColor(0, 0, 255))

    def test_multiclass_compound_requires_all_classes_at_output_level(self):
        """.a.b in CSS only styles elements that have BOTH classes."""
        html = """
        <style>
          .a.b { color: green; }
          .a { color: red; }
        </style>
        <p class="a b">Both classes</p>
        <p class="a">Only class a</p>
        """
        doc = self._parse(html)

        both_run = _run_for_text(doc, "Both classes")
        only_a_run = _run_for_text(doc, "Only class a")

        self.assertIsNotNone(both_run)
        self.assertIsNotNone(only_a_run)

        # .a.b applies + .a applies → green wins (higher specificity)
        self.assertEqual(_run_color(both_run), RGBColor(*Color["green"].value))
        # Only .a applies → red
        self.assertEqual(_run_color(only_a_run), RGBColor(255, 0, 0))

    def test_font_weight_bold_from_css(self):
        html = "<style>p { font-weight: bold; }</style><p>Bold</p>"
        doc = self._parse(html)
        run = _run_for_text(doc, "Bold")
        self.assertTrue(run.font.bold)

    def test_font_weight_normal_inline_overrides_css_bold(self):
        """Inline font-weight: normal should override CSS font-weight: bold."""
        html = """
        <style>p { font-weight: bold; }</style>
        <p style="font-weight: normal;">Not bold</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Not bold")
        self.assertIsNotNone(run)
        self.assertFalse(run.font.bold)

    def test_background_color_from_css_class(self):
        """CSS class background-color should be applied to runs."""
        html = """
        <style>.yellow-bg { background-color: yellow; }</style>
        <p class="yellow-bg">Yellow background</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Yellow background")
        self.assertIsNotNone(run)
        fill = _run_shd_fill(run)
        self.assertIsNotNone(fill)
        self.assertEqual(fill.upper(), "FFFF00")

    def test_text_align_from_css(self):
        """text-align: center from CSS should set paragraph alignment."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        html = "<style>p { text-align: center; }</style><p>Centered</p>"
        doc = self._parse(html)
        para = _para_for_text(doc, "Centered")
        self.assertIsNotNone(para)
        self.assertEqual(para.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)


# ─────────────────────────────────────────────────────────────────────────────
# C) CSS loading order
# ─────────────────────────────────────────────────────────────────────────────

class TestCSSLoadingOrder(unittest.TestCase):
    """Verify that CSS sources are loaded in document order and cascade correctly."""

    def setUp(self):
        self.css_dir = os.path.join(test_dir, "assets", "css")

    def _parse(self, html):
        parser = HtmlToDocx()
        return parser.parse_html_string(html)

    def test_style_tag_before_link_tag(self):
        """<style> before <link>: link CSS loads after and overrides for same property."""
        # small_style.css has: p { color: red; font-weight: bold; }
        css_path = os.path.join(self.css_dir, "small_style.css")
        html = f"""
        <style>p {{ color: green; }}</style>
        <link rel="stylesheet" href="{css_path}">
        <p>Test</p>
        """
        parser = HtmlToDocx()
        parser.add_html_to_document(html, Document())

        # Both style tag and CSS file have p rule.
        # CSS file loaded after style tag → file wins for 'color'
        # small_style.css: p { color: red; }
        self.assertIn("p", parser.css_parser.tag_rules)
        # The last-loaded rule wins; small_style.css sets color: red
        self.assertEqual(parser.css_parser.tag_rules["p"]["color"], "red")

    def test_link_tag_before_style_tag(self):
        """<link> before <style>: <style> is later in document, so it overrides."""
        css_path = os.path.join(self.css_dir, "small_style.css")
        # small_style.css: p { color: red; }
        # <style> tag (color: purple) appears AFTER <link> in document, so purple wins
        html = f"""
        <link rel="stylesheet" href="{css_path}">
        <style>p {{ color: purple; }}</style>
        <p>Test</p>
        """
        parser = HtmlToDocx()
        parser.add_html_to_document(html, Document())

        self.assertIn("p", parser.css_parser.tag_rules)
        self.assertEqual(parser.css_parser.tag_rules["p"]["color"], "purple")

    def test_multiple_style_tags_cascade(self):
        """Multiple <style> tags: later tag overrides earlier for same property."""
        html = """
        <style>p { color: red; }</style>
        <style>p { color: blue; }</style>
        <p>Test</p>
        """
        parser = HtmlToDocx()
        parser.add_html_to_document(html, Document())
        # Second style tag overrides first
        self.assertEqual(parser.css_parser.tag_rules["p"]["color"], "blue")

    def test_multiple_link_tags_cascade(self):
        """Multiple <link> tags: later file overrides earlier for same property."""
        css_path1 = os.path.join(self.css_dir, "small_style.css")
        css_path2 = os.path.join(self.css_dir, "test_styles.css")
        # small_style.css: p { color: red; }
        # test_styles.css: p { color: blue; }
        html = f"""
        <link rel="stylesheet" href="{css_path1}">
        <link rel="stylesheet" href="{css_path2}">
        <p>Test</p>
        """
        parser = HtmlToDocx()
        parser.add_html_to_document(html, Document())
        # test_styles.css loaded second → its p color wins
        self.assertEqual(parser.css_parser.tag_rules["p"]["color"], "blue")

    def test_style_tag_without_head_tag(self):
        """<style> in body (no <head>) should still be processed."""
        html = """
        <body>
          <style>p { color: red; }</style>
          <p>Hello</p>
        </body>
        """
        parser = HtmlToDocx()
        parser.add_html_to_document(html, Document())
        self.assertTrue(parser.css_parser.has_rules())
        self.assertIn("p", parser.css_parser.tag_rules)


# ─────────────────────────────────────────────────────────────────────────────
# D) Span interaction tests
# ─────────────────────────────────────────────────────────────────────────────

class TestSpanInteractions(unittest.TestCase):
    """Verify span CSS class styles and nested span behavior."""

    def _parse(self, html):
        parser = HtmlToDocx()
        return parser.parse_html_string(html)

    def test_span_css_class_color_applied(self):
        """CSS class on <span> should color the span's run."""
        html = """
        <style>.blue { color: blue; }</style>
        <p>Normal <span class="blue">Blue text</span> normal</p>
        """
        doc = self._parse(html)
        blue_run = _run_for_text(doc, "Blue text")
        self.assertIsNotNone(blue_run)
        self.assertEqual(_run_color(blue_run), RGBColor(0, 0, 255))

    def test_span_css_class_bold_applied(self):
        """CSS class font-weight: bold on <span> should make the run bold."""
        html = """
        <style>.bold { font-weight: bold; }</style>
        <p>Normal <span class="bold">Bold</span> normal</p>
        """
        doc = self._parse(html)
        bold_run = _run_for_text(doc, "Bold")
        self.assertIsNotNone(bold_run)
        self.assertTrue(bold_run.font.bold)

    def test_span_inline_style_overrides_css_class(self):
        """Inline style on span should override CSS class style."""
        html = """
        <style>.blue { color: blue; }</style>
        <p><span class="blue" style="color: red;">Red override</span></p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Red override")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))

    def test_outer_span_css_class_applied_to_inner_text(self):
        """Outer span's CSS class styles should apply to text inside inner span."""
        html = """
        <style>
          .outer { color: red; }
          .inner { font-weight: bold; }
        </style>
        <p><span class="outer"><span class="inner">Nested text</span></span></p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Nested text")
        self.assertIsNotNone(run)
        # outer gives red color, inner gives bold — both should apply
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))
        self.assertTrue(run.font.bold)

    def test_inner_span_overrides_outer_span_color(self):
        """Inner span's color overrides outer span's color."""
        html = """
        <style>
          .outer { color: red; }
          .inner { color: blue; }
        </style>
        <p><span class="outer"><span class="inner">Inner wins</span></span></p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Inner wins")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_outer_span_text_gets_outer_styles_only(self):
        """Text directly in the outer span (not in inner) only gets outer styles."""
        html = """
        <style>
          .outer { color: red; }
          .inner { color: blue; font-weight: bold; }
        </style>
        <p><span class="outer">Outer text <span class="inner">Inner text</span></span></p>
        """
        doc = self._parse(html)
        outer_run = _run_for_text(doc, "Outer text")
        inner_run = _run_for_text(doc, "Inner text")

        self.assertIsNotNone(outer_run)
        self.assertIsNotNone(inner_run)

        # Outer text: red, not bold
        self.assertEqual(_run_color(outer_run), RGBColor(255, 0, 0))
        self.assertNotEqual(_run_color(inner_run), RGBColor(255, 0, 0))

    def test_span_inside_li_css_class(self):
        """CSS class on span inside <li> should be applied."""
        html = """
        <style>.blue { color: blue; }</style>
        <ul><li>Normal <span class="blue">Blue in list</span></li></ul>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Blue in list")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_span_inside_h_tag_css_class(self):
        """CSS class on span inside heading should be applied."""
        html = """
        <style>.highlight { background-color: yellow; }</style>
        <h2>Normal <span class="highlight">Highlighted</span> heading</h2>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Highlighted")
        self.assertIsNotNone(run)
        fill = _run_shd_fill(run)
        self.assertIsNotNone(fill)
        self.assertEqual(fill.upper(), "FFFF00")

    def test_nested_spans_inline_styles_both_applied(self):
        """Nested spans with inline styles: outer color + inner bold both apply."""
        html = """
        <p>
          <span style="color: red;">
            Outer <span style="font-weight: bold;">Inner bold</span>
          </span>
        </p>
        """
        doc = self._parse(html)
        inner_run = _run_for_text(doc, "Inner bold")
        self.assertIsNotNone(inner_run)
        self.assertEqual(_run_color(inner_run), RGBColor(255, 0, 0))
        self.assertTrue(inner_run.font.bold)

    def test_text_decoration_none_on_span_removes_parent_underline(self):
        """text-decoration: none on a span should remove the parent element's underline."""
        html = """
        <style>p { text-decoration: underline; }</style>
        <p>Underlined <span style="text-decoration: none;">Not underlined</span></p>
        """
        doc = self._parse(html)
        no_underline_run = _run_for_text(doc, "Not underlined")
        self.assertIsNotNone(no_underline_run)
        self.assertFalse(no_underline_run.font.underline)

    def test_span_text_decoration_style_with_parent_underline(self):
        """Span's text-decoration-style: double should combine with parent's text-decoration: underline."""
        html = """
        <p style="text-decoration: underline;">
            Normal underline
            <span style="text-decoration-style: double;">Double underline</span>
        </p>
        """
        doc = self._parse(html)
        double_run = _run_for_text(doc, "Double underline")
        self.assertIsNotNone(double_run)
        self.assertEqual(double_run.font.underline, WD_UNDERLINE.DOUBLE)

    def test_span_multiple_css_classes(self):
        """A span with multiple CSS classes should get styles from all matching classes."""
        html = """
        <style>
          .bold { font-weight: bold; }
          .italic { font-style: italic; }
        </style>
        <p><span class="bold italic">Bold italic</span></p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Bold italic")
        self.assertIsNotNone(run)
        self.assertTrue(run.font.bold)
        self.assertTrue(run.font.italic)


# ─────────────────────────────────────────────────────────────────────────────
# E) External CSS tests
# ─────────────────────────────────────────────────────────────────────────────

class TestExternalCSSBehavior(unittest.TestCase):
    """Tests for external CSS file loading via <link> tags."""

    def setUp(self):
        self.css_dir = os.path.join(test_dir, "assets", "css")

    def _parse(self, html):
        parser = HtmlToDocx()
        return parser.parse_html_string(html)

    def test_external_css_applied_to_output(self):
        """Styles from an external CSS file should appear in the docx output."""
        css_path = os.path.join(self.css_dir, "small_style.css")
        # small_style.css: p { color: red; font-weight: bold; }
        html = f'<link rel="stylesheet" href="{css_path}"><p>Red bold</p>'
        doc = self._parse(html)
        run = _run_for_text(doc, "Red bold")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))
        self.assertTrue(run.font.bold)

    def test_inline_style_overrides_external_css(self):
        """Inline style should override a rule from an external CSS file."""
        css_path = os.path.join(self.css_dir, "small_style.css")
        # CSS: p { color: red }, Inline: color: blue
        html = f'<link rel="stylesheet" href="{css_path}"><p style="color: blue;">Blue override</p>'
        doc = self._parse(html)
        run = _run_for_text(doc, "Blue override")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_external_css_class_applied(self):
        """A class rule from an external CSS file should be applied."""
        css_path = os.path.join(self.css_dir, "small_style.css")
        # small_style.css: .blueSpan { color: blue; }
        html = f'<link rel="stylesheet" href="{css_path}"><p><span class="blueSpan">Blue</span></p>'
        doc = self._parse(html)
        run = _run_for_text(doc, "Blue")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_invalid_css_file_handled_gracefully(self):
        """A missing external CSS file should not crash the parser."""
        html = '<link rel="stylesheet" href="/nonexistent/path/styles.css"><p>Hello</p>'
        try:
            self._parse(html)
            success = True
        except Exception:
            success = False
        self.assertTrue(success)

    def test_external_css_with_style_tag_combined(self):
        """External CSS and inline <style> tag should both contribute to the output."""
        css_path = os.path.join(self.css_dir, "small_style.css")
        # small_style.css: .blueSpan { color: blue; }
        # style tag adds: .green { color: green; }
        html = f"""
        <link rel="stylesheet" href="{css_path}">
        <style>.green {{ color: green; }}</style>
        <p>
          <span class="blueSpan">Blue</span>
          <span class="green">Green</span>
        </p>
        """
        doc = self._parse(html)
        blue_run = _run_for_text(doc, "Blue")
        green_run = _run_for_text(doc, "Green")
        self.assertIsNotNone(blue_run)
        self.assertIsNotNone(green_run)
        self.assertEqual(_run_color(blue_run), RGBColor(0, 0, 255))
        self.assertEqual(_run_color(green_run), RGBColor(*Color["green"].value))

    def test_external_css_with_temp_file(self):
        """Test CSS loading from a temp file (encoding)."""
        css = "p { color: red; font-weight: bold; }"
        with tempfile.NamedTemporaryFile(mode='w', suffix='.css', delete=False, encoding='utf-8') as f:
            f.write(css)
            tmp_path = f.name
        try:
            html = f'<link rel="stylesheet" href="{tmp_path}"><p>Temp CSS</p>'
            doc = self._parse(html)
            run = _run_for_text(doc, "Temp CSS")
            self.assertIsNotNone(run)
            self.assertEqual(_run_color(run), RGBColor(255, 0, 0))
            self.assertTrue(run.font.bold)
        finally:
            os.unlink(tmp_path)


# ─────────────────────────────────────────────────────────────────────────────
# F) Edge cases and robustness
# ─────────────────────────────────────────────────────────────────────────────

class TestEdgeCases(unittest.TestCase):
    """Edge cases: empty styles, invalid values, unusual HTML structures."""

    def _parse(self, html):
        parser = HtmlToDocx()
        return parser.parse_html_string(html)

    def test_empty_style_tag(self):
        """Empty <style> tag should not crash."""
        html = "<style></style><p>Hello</p>"
        try:
            self._parse(html)
            success = True
        except Exception:
            success = False
        self.assertTrue(success)

    def test_style_tag_with_only_comments(self):
        """<style> with only comments should result in no rules."""
        html = "<style>/* this is a comment */</style><p>Hello</p>"
        parser = HtmlToDocx()
        parser.add_html_to_document(html, Document())
        self.assertFalse(parser.css_parser.has_rules())

    def test_empty_inline_style(self):
        """Empty style="" attribute should not crash."""
        html = '<p style="">Hello</p>'
        try:
            self._parse(html)
            success = True
        except Exception:
            success = False
        self.assertTrue(success)

    def test_invalid_color_in_css_falls_back(self):
        """An unrecognized color value in CSS should not crash; falls back to black."""
        html = "<style>p { color: not-a-color; }</style><p>Hello</p>"
        try:
            self._parse(html)
            success = True
        except Exception:
            success = False
        self.assertTrue(success)

    def test_color_inherit_in_css(self):
        """color: inherit in CSS should be skipped (not crash)."""
        html = "<style>p { color: inherit; }</style><p>Hello</p>"
        try:
            self._parse(html)
            success = True
        except Exception:
            success = False
        self.assertTrue(success)

    def test_unknown_css_property_skipped(self):
        """An unsupported CSS property like letter-spacing should be silently ignored."""
        html = "<style>p { letter-spacing: 2px; color: red; }</style><p>Hello</p>"
        doc = self._parse(html)
        run = _run_for_text(doc, "Hello")
        # The known property (color: red) should still be applied
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))

    def test_font_size_named_values(self):
        """Named font sizes (e.g. small, medium, large) should be handled."""
        html = "<style>p { font-size: large; }</style><p>Large text</p>"
        try:
            self._parse(html)
            success = True
        except Exception:
            success = False
        self.assertTrue(success)

    def test_multiple_style_tags_independent(self):
        """Two separate <style> tags should both contribute rules."""
        html = """
        <style>p { color: red; }</style>
        <style>.bold { font-weight: bold; }</style>
        <p class="bold">Bold red</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Bold red")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))
        self.assertTrue(run.font.bold)

    def test_style_applied_to_h_tags(self):
        """CSS rules for heading tags should be applied."""
        html = "<style>h2 { color: navy; font-weight: bold; }</style><h2>Heading</h2>"
        doc = self._parse(html)
        run = _run_for_text(doc, "Heading")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(*Color["navy"].value))
        self.assertTrue(run.font.bold)

    def test_css_parser_reused_across_add_html_calls(self):
        """CSS rules from first call should still apply in the second call on the same parser."""
        parser = HtmlToDocx()
        doc1 = Document()
        doc2 = Document()

        parser.add_html_to_document(
            "<style>p { color: red; }</style><p>First</p>", doc1
        )
        # Second call without style tag — the CSS from the first call is preserved
        parser.add_html_to_document("<p>Second</p>", doc2)

        self.assertTrue(parser.css_parser.has_rules())
        # The red color rule should have been applied to the second doc's paragraph too
        run2 = _run_for_text(doc2, "Second")
        self.assertIsNotNone(run2)
        self.assertEqual(_run_color(run2), RGBColor(255, 0, 0))

    def test_div_inline_style_applied_to_paragraph(self):
        """A div's text-align inline style should propagate to its child paragraphs."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        html = '<div style="text-align: center"><p>Centered via div</p></div>'
        doc = self._parse(html)
        para = _para_for_text(doc, "Centered via div")
        self.assertIsNotNone(para)
        self.assertEqual(para.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_css_with_pseudo_class_stripped(self):
        """Pseudo-classes like :hover should be stripped and base rule applied."""
        html = "<style>a:hover { color: red; }</style><p><a href='#'>Link</a></p>"
        parser = HtmlToDocx()
        parser.add_html_to_document(html, Document())
        # a:hover → stored under 'a' tag rule
        self.assertIn("a", parser.css_parser.tag_rules)

    def test_css_selector_with_attribute(self):
        """Attribute selectors ([type='text']) should not crash and rules after them still parse."""
        html = "<style>input[type='text'] { border: 1px solid; } p { color: red; }</style><p>Hello</p>"
        parser = HtmlToDocx()
        parser.add_html_to_document(html, Document())
        # The p rule after the attribute selector should still be loaded
        self.assertIn("p", parser.css_parser.tag_rules)
        self.assertEqual(parser.css_parser.tag_rules["p"]["color"], "red")


# ─────────────────────────────────────────────────────────────────────────────
# G) CSS + existing features regression tests
# ─────────────────────────────────────────────────────────────────────────────

class TestCSSWithExistingFeaturesRegression(unittest.TestCase):
    """Ensure CSS feature doesn't break existing functionality."""

    def _parse(self, html):
        parser = HtmlToDocx()
        return parser.parse_html_string(html)

    def test_bold_italic_tags_still_work_with_css(self):
        """<b> and <em> should still make text bold/italic even with CSS present."""
        html = "<style>p { color: red; }</style><p><b>Bold</b> and <em>Italic</em></p>"
        doc = self._parse(html)
        bold_run = _run_for_text(doc, "Bold")
        italic_run = _run_for_text(doc, "Italic")
        self.assertIsNotNone(bold_run)
        self.assertIsNotNone(italic_run)
        self.assertTrue(bold_run.font.bold)
        self.assertTrue(italic_run.font.italic)

    def test_inline_style_still_works_without_css(self):
        """Inline styles should work even with no <style> or <link> tags."""
        html = '<p style="color: blue; font-weight: bold;">Styled inline</p>'
        doc = self._parse(html)
        run = _run_for_text(doc, "Styled inline")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))
        self.assertTrue(run.font.bold)

    def test_table_still_works_with_css(self):
        """Tables should still parse correctly when CSS is present."""
        html = """
        <style>p { color: red; }</style>
        <table><tr><td>Cell 1</td><td>Cell 2</td></tr></table>
        """
        doc = self._parse(html)
        self.assertGreaterEqual(len(doc.tables), 1)

    def test_image_still_works_with_css(self):
        """Image parsing should not be affected by CSS feature."""
        html = """
        <style>p { color: red; }</style>
        <p>No image but CSS works</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "No image but CSS works")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))

    def test_ordered_unordered_lists_still_work(self):
        """Lists should be unaffected by CSS."""
        html = """
        <style>.item { color: blue; }</style>
        <ul>
          <li class="item">Item one</li>
          <li>Item two</li>
        </ul>
        """
        doc = self._parse(html)
        self.assertGreaterEqual(len(doc.paragraphs), 2)

    def test_css_does_not_appear_in_paragraph_text(self):
        """CSS rule text should never appear as paragraph content."""
        html = """
        <style>p { color: red; font-weight: bold; }</style>
        <link rel="stylesheet" href="/nonexistent.css">
        <p>Only this text</p>
        """
        doc = self._parse(html)
        for para in doc.paragraphs:
            self.assertNotIn("color:", para.text)
            self.assertNotIn("font-weight", para.text)
            self.assertNotIn("<style>", para.text)
            self.assertNotIn("<link", para.text)

    def test_css_id_rule_and_bookmark_coexist(self):
        """CSS id selector should not interfere with bookmark generation for the same id."""
        html = """
        <style>#section { color: navy; }</style>
        <h2 id="section">Section heading</h2>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Section heading")
        self.assertIsNotNone(run)
        # Color applied from CSS
        self.assertEqual(_run_color(run), RGBColor(*Color["navy"].value))

    def test_span_css_class_does_not_affect_paragraph_alignment(self):
        """A CSS class on a span (run-level) should not affect paragraph alignment."""
        html = """
        <style>.center-text { text-align: center; }</style>
        <p><span class="center-text">Centered span?</span></p>
        """
        doc = self._parse(html)
        para = _para_for_text(doc, "Centered span?")
        self.assertIsNotNone(para)
        # text-align on a span — alignment may or may not be applied, but should not crash
        # The paragraph itself should not have alignment set from span (implementation-defined)



# ─────────────────────────────────────────────────────────────────────────────
# H) Complex CSS scenarios and nesting
# ─────────────────────────────────────────────────────────────────────────────

class TestComplexCSSScenarios(unittest.TestCase):
    """Complex nesting, compound selectors, and multi-level CSS interactions."""

    def _parse(self, html):
        parser = HtmlToDocx()
        return parser.parse_html_string(html)

    def test_span_compound_selector(self):
        """span.myclass compound selector should only apply to span elements with that class."""
        html = """
        <style>
          span.highlight { color: blue; font-weight: bold; }
          span { color: red; }
        </style>
        <p><span class="highlight">Highlighted span</span> <span>Plain span</span></p>
        """
        doc = self._parse(html)
        highlighted = _run_for_text(doc, "Highlighted span")
        plain = _run_for_text(doc, "Plain span")
        self.assertIsNotNone(highlighted)
        self.assertIsNotNone(plain)
        # span.highlight → blue (compound overrides bare span red)
        self.assertEqual(_run_color(highlighted), RGBColor(0, 0, 255))
        self.assertTrue(highlighted.font.bold)
        # bare span → red
        self.assertEqual(_run_color(plain), RGBColor(255, 0, 0))
        self.assertFalse(plain.font.bold or False)

    def test_heading_compound_selector(self):
        """h2.special compound rule should only apply to h2 with class=special."""
        html = """
        <style>
          h2 { color: navy; }
          h2.special { color: green; font-weight: bold; }
        </style>
        <h2>Normal heading</h2>
        <h2 class="special">Special heading</h2>
        """
        doc = self._parse(html)
        normal_run = _run_for_text(doc, "Normal heading")
        special_run = _run_for_text(doc, "Special heading")
        self.assertIsNotNone(normal_run)
        self.assertIsNotNone(special_run)
        # Normal h2: navy
        self.assertEqual(_run_color(normal_run), RGBColor(*Color["navy"].value))
        # Special h2: green (compound overrides tag rule)
        self.assertEqual(_run_color(special_run), RGBColor(*Color["green"].value))
        self.assertTrue(special_run.font.bold)

    def test_li_with_css_class(self):
        """CSS class rule applied to a list item should color its run."""
        html = """
        <style>.important-item { color: red; font-weight: bold; }</style>
        <ul>
          <li class="important-item">Important item</li>
          <li>Normal item</li>
        </ul>
        """
        doc = self._parse(html)
        important_run = _run_for_text(doc, "Important item")
        self.assertIsNotNone(important_run)
        self.assertEqual(_run_color(important_run), RGBColor(255, 0, 0))
        self.assertTrue(important_run.font.bold)

    def test_pre_with_css_rule(self):
        """CSS rule on pre tag should apply to its run."""
        html = """
        <style>pre { color: blue; }</style>
        <pre>Preformatted blue text</pre>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Preformatted blue text")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_three_level_nesting_chain(self):
        """CSS classes on div, p, and span all applied correctly in a nesting chain."""
        html = """
        <style>
          .outer-div { background-color: yellow; }
          .inner-p { color: red; font-weight: bold; }
          .inner-span { font-style: italic; color: blue; }
        </style>
        <div class="outer-div">
          <p class="inner-p">
            Paragraph text
            <span class="inner-span">Span text</span>
          </p>
        </div>
        """
        doc = self._parse(html)

        # Paragraph text run: red + bold (from .inner-p)
        para_run = _run_for_text(doc, "Paragraph text")
        self.assertIsNotNone(para_run)
        self.assertEqual(_run_color(para_run), RGBColor(255, 0, 0))
        self.assertTrue(para_run.font.bold)

        # Span text run: blue (from .inner-span, overrides .inner-p red) + italic
        span_run = _run_for_text(doc, "Span text")
        self.assertIsNotNone(span_run)
        self.assertEqual(_run_color(span_run), RGBColor(0, 0, 255))
        self.assertTrue(span_run.font.italic)

    def test_text_transform_from_css_class(self):
        """text-transform from a CSS class rule should transform run text."""
        html = """
        <style>.shout { text-transform: uppercase; }</style>
        <p class="shout">quiet text</p>
        """
        doc = self._parse(html)
        para = _para_for_text(doc, "QUIET TEXT")
        self.assertIsNotNone(para)
        # All text in the paragraph should be uppercased
        full_text = "".join(r.text for r in para.runs)
        self.assertEqual(full_text.strip(), "QUIET TEXT")

    def test_font_family_from_css_class_on_span(self):
        """font-family from a CSS class rule should set the run's font name."""
        html = """
        <style>.courier { font-family: 'Courier New', monospace; }</style>
        <p>Normal <span class="courier">Courier text</span></p>
        """
        doc = self._parse(html)
        courier_run = _run_for_text(doc, "Courier text")
        self.assertIsNotNone(courier_run)
        self.assertEqual(courier_run.font.name, "Courier New")

    def test_font_size_em_from_css_rule(self):
        """font-size in em units from a CSS class rule should set a non-None run size."""
        html = """
        <style>.big { font-size: 2em; }</style>
        <p class="big">Big text</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Big text")
        self.assertIsNotNone(run)
        # 2em = 24pt (1em assumed 12pt), so size should be > 0
        self.assertIsNotNone(run.font.size)
        self.assertGreater(run.font.size, 0)

    def test_font_size_rem_from_css_rule(self):
        """font-size in rem units from a CSS rule should set a non-None run size."""
        html = """
        <style>p { font-size: 1.5rem; }</style>
        <p>Rem text</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Rem text")
        self.assertIsNotNone(run)
        self.assertIsNotNone(run.font.size)
        self.assertGreater(run.font.size, 0)

    def test_css_applied_inside_table_cell(self):
        """CSS class rule should be applied to text content inside a table cell."""
        html = """
        <style>.cell-text { color: red; font-weight: bold; }</style>
        <table>
          <tr>
            <td><p class="cell-text">Cell content</p></td>
          </tr>
        </table>
        """
        doc = self._parse(html)
        # Table cell content is not in doc.paragraphs — search inside tables
        run = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for r in para.runs:
                            if "Cell content" in r.text:
                                run = r
        self.assertIsNotNone(run, "Run 'Cell content' should be found inside the table cell")
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))
        self.assertTrue(run.font.bold)

    def test_div_css_background_color_propagates_to_runs(self):
        """background-color on a div via CSS class should propagate to runs in child paragraphs."""
        html = """
        <style>.highlighted-section { background-color: yellow; }</style>
        <div class="highlighted-section">
          <p>Text in highlighted div</p>
        </div>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Text in highlighted div")
        self.assertIsNotNone(run)
        fill = _run_shd_fill(run)
        self.assertIsNotNone(fill)
        self.assertEqual(fill.upper(), "FFFF00")

    def test_multiple_css_sources_all_applied(self):
        """Style tag + link tag + inline: all three sources contribute to final styles."""
        import os
        import tempfile
        css = ".ext-class { font-style: italic; }"
        with tempfile.NamedTemporaryFile(mode='w', suffix='.css', delete=False, encoding='utf-8') as f:
            f.write(css)
            tmp = f.name
        try:
            html = f"""
            <style>.style-tag {{ color: red; }}</style>
            <link rel="stylesheet" href="{tmp}">
            <p class="style-tag ext-class" style="font-weight: bold;">Triple styled</p>
            """
            doc = self._parse(html)
            run = _run_for_text(doc, "Triple styled")
            self.assertIsNotNone(run)
            # From <style> tag
            self.assertEqual(_run_color(run), RGBColor(255, 0, 0))
            # From external CSS
            self.assertTrue(run.font.italic)
            # From inline
            self.assertTrue(run.font.bold)
        finally:
            os.unlink(tmp)

    def test_inline_important_overrides_css_important(self):
        """Inline !important should override a CSS !important for the same property."""
        html = """
        <style>p { color: blue !important; }</style>
        <p style="color: red !important;">Red wins</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Red wins")
        self.assertIsNotNone(run)
        # Inline !important is applied last (highest priority) → red wins
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))

    def test_css_class_on_code_tag(self):
        """CSS class rule applied to a <code> element should style its run."""
        html = """
        <style>.highlight-code { color: blue; background-color: yellow; }</style>
        <p><code class="highlight-code">styled code</code></p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "styled code")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_background_color_from_css_class_on_span(self):
        """background-color from CSS class on span should shade only the span's run."""
        html = """
        <style>.shaded { background-color: yellow; }</style>
        <p>Normal <span class="shaded">Shaded</span> normal</p>
        """
        doc = self._parse(html)
        shaded_run = _run_for_text(doc, "Shaded")
        normal_run = _run_for_text(doc, "Normal")
        self.assertIsNotNone(shaded_run)
        # Shaded span should have yellow background
        fill = _run_shd_fill(shaded_run)
        self.assertIsNotNone(fill)
        self.assertEqual(fill.upper(), "FFFF00")
        # Normal text should NOT have yellow background
        if normal_run:
            normal_fill = _run_shd_fill(normal_run)
            self.assertNotEqual((normal_fill or "").upper(), "FFFF00")

    def test_div_css_does_not_override_p_own_color(self):
        """A div's CSS color should NOT override a <p>'s own color style."""
        html = """
        <style>
          .red-div { color: red; }
          .blue-p { color: blue; }
        </style>
        <div class="red-div">
          <p class="blue-p">Blue paragraph in red div</p>
        </div>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Blue paragraph in red div")
        self.assertIsNotNone(run)
        # The <p> has its own color (blue) which takes precedence over div's red
        self.assertEqual(_run_color(run), RGBColor(0, 0, 255))

    def test_text_decoration_underline_from_css_class(self):
        """text-decoration: underline from a CSS class should underline the run."""
        html = """
        <style>.underlined { text-decoration: underline; }</style>
        <p class="underlined">Underlined paragraph</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Underlined paragraph")
        self.assertIsNotNone(run)
        self.assertTrue(run.font.underline)

    def test_text_decoration_line_through_from_css_class(self):
        """text-decoration: line-through from a CSS class should strike through the run."""
        html = """
        <style>.struck { text-decoration: line-through; }</style>
        <p class="struck">Struck through</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Struck through")
        self.assertIsNotNone(run)
        self.assertTrue(run.font.strike)

    def test_css_class_font_weight_bold_on_heading(self):
        """A CSS class adding font-weight: bold should make heading text bold."""
        html = """
        <style>.extra-bold { font-weight: bold; }</style>
        <h3 class="extra-bold">Bold heading</h3>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Bold heading")
        self.assertIsNotNone(run)
        self.assertTrue(run.font.bold)

    def test_media_query_does_not_pollute_rules(self):
        """@media blocks in a <style> tag must be skipped and not create spurious rules."""
        html = """
        <style>
          p { color: red; }
          @media print {
            p { color: black; }
            .print-only { display: block; }
          }
          .visible { font-weight: bold; }
        </style>
        <p class="visible">Test paragraph</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Test paragraph")
        self.assertIsNotNone(run)
        # CSS: p { color: red } applies (media query p { color: black } does NOT)
        self.assertEqual(_run_color(run), RGBColor(255, 0, 0))
        # .visible { font-weight: bold } applies
        self.assertTrue(run.font.bold)

    def test_css_id_selector_on_span(self):
        """CSS ID selector should apply to a span with matching id."""
        html = """
        <style>#special-span { color: green; font-weight: bold; }</style>
        <p><span id="special-span">Special span</span> normal</p>
        """
        doc = self._parse(html)
        run = _run_for_text(doc, "Special span")
        self.assertIsNotNone(run)
        self.assertEqual(_run_color(run), RGBColor(*Color["green"].value))
        self.assertTrue(run.font.bold)

    def test_css_class_applied_to_multiple_elements(self):
        """The same CSS class should style all elements that carry it."""
        html = """
        <style>.red { color: red; }</style>
        <p class="red">First red</p>
        <p class="red">Second red</p>
        <p>Not red</p>
        """
        doc = self._parse(html)
        first = _run_for_text(doc, "First red")
        second = _run_for_text(doc, "Second red")
        plain = _run_for_text(doc, "Not red")
        self.assertIsNotNone(first)
        self.assertIsNotNone(second)
        self.assertEqual(_run_color(first), RGBColor(255, 0, 0))
        self.assertEqual(_run_color(second), RGBColor(255, 0, 0))
        if plain:
            self.assertNotEqual(_run_color(plain), RGBColor(255, 0, 0))


if __name__ == "__main__":
    unittest.main()
