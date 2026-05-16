"""Document extraction — turn binary files into LLM-readable text.

Sage's READ tool reads files as UTF-8 text. PDFs, Word docs, spreadsheets,
etc. all decode as garbage that bloats the model's context and produces
hallucinated answers.

This extractor sniffs the file format (magic bytes + extension) and routes
through the right parser:

    PDF       → pypdf
    DOCX      → python-docx
    Plain     → utf-8 read (with safe-decode fallback)

Parsers are *optional* deps: if pypdf isn't installed, we raise a clear
UnsupportedFormatError telling the user how to enable PDF support, rather
than crashing or returning bytes-as-text.

CLI integration: the READ tool auto-detects and routes; users don't have
to specify the format.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger("sage.document_extractor")


# Probe optional deps at import time so we know what we can handle.
# Tests can monkey-patch these flags to simulate missing deps.
try:
    import pypdf  # noqa: F401
    _HAS_PYPDF = True
except Exception:
    _HAS_PYPDF = False

try:
    import docx  # python-docx  # noqa: F401
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False


class UnsupportedFormatError(RuntimeError):
    """Raised when we can't parse a file — either format unknown or the
    required optional dep isn't installed. Always carries an actionable
    message about how to fix it."""


# ── Data type ────────────────────────────────────────────────────────────────


@dataclass(frozen=True)
class ExtractedDocument:
    """Parsed contents of a document. ``text`` is the LLM-ready content;
    ``page_count`` helps the caller decide whether to chunk."""
    source_path: Path
    text: str
    page_count: int
    mime_type: str


# ── Format detection ─────────────────────────────────────────────────────────


# Magic-byte prefixes for known formats. Order matters: PDF check comes
# before generic ZIP because DOCX is a ZIP-wrapped XML bundle and would
# otherwise mis-classify.
_MAGIC_PREFIXES: tuple[tuple[bytes, str], ...] = (
    (b"%PDF-", "pdf"),
    (b"PK\x03\x04", "docx"),  # Also XLSX/PPTX — we narrow by extension later
)


# Extensions we trust as text (no magic-byte signature needed; common
# plain-text file types and code/config formats).
_TEXT_EXTENSIONS: frozenset[str] = frozenset({
    ".txt", ".md", ".markdown", ".rst",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".go", ".rs", ".rb",
    ".c", ".cpp", ".h", ".hpp", ".cs", ".sh", ".bash", ".zsh",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".html", ".htm", ".css", ".scss", ".sql", ".graphql",
    ".csv", ".tsv", ".log", ".env",
})


class DocumentExtractor:
    """Sniffs file format and dispatches to the right parser."""

    @staticmethod
    def detect_format(path: Path) -> str:
        """Return one of: 'pdf', 'docx', 'text'. Raises if file missing."""
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        # 1. Magic bytes — most reliable signal
        with path.open("rb") as f:
            head = f.read(8)
        for prefix, fmt in _MAGIC_PREFIXES:
            if head.startswith(prefix):
                # DOCX/XLSX/PPTX all share the ZIP magic. Use extension to
                # disambiguate; default to "docx" if user named it as such.
                if fmt == "docx":
                    ext = path.suffix.lower()
                    if ext == ".docx":
                        return "docx"
                    # Unknown ZIP-wrapped format — bubble up as unsupported.
                    # (Future: route .xlsx through openpyxl.)
                    raise UnsupportedFormatError(
                        f"{path.name} looks like a ZIP archive but not a "
                        f"recognized document format (got extension {ext!r}). "
                        "Supported: .docx"
                    )
                return fmt

        # 2. Known text extensions
        if path.suffix.lower() in _TEXT_EXTENSIONS:
            return "text"

        # 3. Heuristic: does it look like ASCII text?
        try:
            head_text = head.decode("utf-8", errors="strict")
            # All printable + whitespace = plausibly text
            if all(c.isprintable() or c in "\n\r\t " for c in head_text):
                return "text"
        except UnicodeDecodeError:
            pass

        raise UnsupportedFormatError(
            f"Could not identify format of {path.name}. "
            f"Supported: PDF, DOCX, plain text (txt/md/code/config). "
            f"Got first bytes: {head!r}"
        )

    def extract(self, path: str | Path) -> ExtractedDocument:
        """Parse a document and return its text + metadata."""
        path = Path(path)
        fmt = self.detect_format(path)
        if fmt == "pdf":
            return self._extract_pdf(path)
        if fmt == "docx":
            return self._extract_docx(path)
        if fmt == "text":
            return self._extract_text(path)
        # Defensive — detect_format should have raised
        raise UnsupportedFormatError(f"Unknown format: {fmt}")

    # ── Format-specific extractors ────────────────────────────────────────────

    def _extract_pdf(self, path: Path) -> ExtractedDocument:
        if not _HAS_PYPDF:
            raise UnsupportedFormatError(
                f"PDF support requires pypdf. Install with: pip install pypdf\n"
                f"Could not read: {path}"
            )
        import pypdf

        reader = pypdf.PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception as exc:
                # Some PDFs have malformed pages — skip the page, don't fail
                # the whole document. Log so we can debug recurring issues.
                logger.warning("Skipping unreadable PDF page in %s: %s", path, exc)
                pages.append("")

        text = _normalize_whitespace("\n\n".join(pages))
        return ExtractedDocument(
            source_path=path,
            text=text,
            page_count=len(reader.pages),
            mime_type="application/pdf",
        )

    def _extract_docx(self, path: Path) -> ExtractedDocument:
        if not _HAS_DOCX:
            raise UnsupportedFormatError(
                f"DOCX support requires python-docx. Install with: pip install python-docx\n"
                f"Could not read: {path}"
            )
        import docx

        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        # Tables aren't iterated by .paragraphs — pull them separately.
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells if c.text]
                if cells:
                    paragraphs.append(" | ".join(cells))

        text = _normalize_whitespace("\n\n".join(paragraphs))
        return ExtractedDocument(
            source_path=path,
            text=text,
            page_count=1,  # DOCX has no fixed page count without rendering
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def _extract_text(self, path: Path) -> ExtractedDocument:
        # Use errors='replace' so a stray non-utf8 byte in a code file
        # doesn't blow up extraction. Same behavior as sage's existing
        # READ tool.
        text = path.read_text(encoding="utf-8", errors="replace")
        return ExtractedDocument(
            source_path=path,
            text=text,
            page_count=1,
            mime_type="text/plain",
        )


# ── Helpers ──────────────────────────────────────────────────────────────────


_MULTI_BLANK = re.compile(r"\n{3,}")


def _normalize_whitespace(text: str) -> str:
    """Trim trailing whitespace per line and collapse 3+ blank lines to 2.
    PDFs produce a lot of noise; this keeps the LLM context lean without
    losing structural breaks."""
    lines = [line.rstrip() for line in text.splitlines()]
    return _MULTI_BLANK.sub("\n\n", "\n".join(lines)).strip()


__all__ = [
    "DocumentExtractor",
    "ExtractedDocument",
    "UnsupportedFormatError",
]
